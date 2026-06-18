# -*- coding: utf-8 -*-
"""
DOCX 读写工具函数库
===================
提供两组功能：
  [写入] 精确控制 DOCX 格式（来自 latex2docx.py 的经验积累）
  [读取] 解析已有 DOCX 的格式信息（用于模板分析）
"""
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
import io, os, json, re

# ======================================================================
# 第一部分：DOCX 写入工具（来源于 latex2docx.py 的实战经验）
# ======================================================================

def set_xml_spacing(p, line_val, is_exact=True):
    """设置段落行距（XML 层面精确控制）"""
    pPr = p._element.get_or_add_pPr()
    sp = pPr.find(qn('w:spacing'))
    if sp is None:
        sp = OxmlElement('w:spacing')
        pPr.append(sp)
    sp.set(qn('w:line'), str(line_val))
    sp.set(qn('w:lineRule'), 'exact' if is_exact else 'auto')


def set_xml_indent(p, first=None, left=None, right=None):
    """设置段落缩进（单位：pt，内部转为 twips）"""
    pPr = p._element.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    if first is not None:
        ind.set(qn('w:firstLine'), str(int(first * 20)))
    if left is not None:
        ind.set(qn('w:left'), str(int(left * 20)))
    if right is not None:
        ind.set(qn('w:right'), str(int(right * 20)))


def set_xml_font(r, cn='宋体', en=None, sz=12, bold=False, hint=True, underline=None):
    """设置 run 字体（eastAsia 中文字体 + hint 属性）
       underline: None=不设置, 'single'=单下划线, 'double'=双下划线
    """
    r.font.size = Pt(sz)
    r.font.bold = bold
    if underline:
        r.font.underline = True
    rPr = r._element.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rPr.insert(0, rf)
    rf.set(qn('w:eastAsia'), cn)
    if en:
        rf.set(qn('w:ascii'), en)
        rf.set(qn('w:hAnsi'), en)
    if hint:
        rf.set(qn('w:hint'), 'eastAsia')


def load_img_bytes(path, max_w=1200):
    """
    加载图片，缩放到指定宽度，返回 JPEG 字节流。
    path: 图片文件路径
    max_w: 最大宽度像素（默认 1200）
    """
    if not path or not os.path.exists(path):
        return None
    img = Image.open(path)
    if img.size[0] > max_w:
        ratio = max_w / img.size[0]
        new_h = int(img.size[1] * ratio)
        img = img.resize((max_w, new_h), Image.LANCZOS)
    if img.mode == 'RGBA':
        bg = Image.new('RGB', img.size, (255, 255, 255))
        bg.paste(img, mask=img.split()[3])
        img = bg
    buf = io.BytesIO()
    img.save(buf, 'JPEG', quality=92, optimize=True)
    buf.seek(0)
    return buf


def cv_line(doc, text, pt, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
            li=None, ri=None, fi=None, cn='宋体', en=None):
    """封面段落生成器：独立格式，1.5 倍行距"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    set_xml_spacing(p, 360, False)
    if fi:
        set_xml_indent(p, first=fi)
    if li:
        set_xml_indent(p, left=li)
    if ri:
        set_xml_indent(p, right=ri)
    r = p.add_run(text)
    set_xml_font(r, cn, en, sz=pt, bold=bold)
    return p


def cv_empty(doc):
    """封面空行（居中，1.5 倍行距）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    set_xml_spacing(p, 360, False)
    return p


# ======================================================================
# 第二部分：DOCX 读取工具（用于模板分析）
# ======================================================================

def _get_val(element, attr, default=None):
    """安全读取 XML 属性"""
    if element is None:
        return default
    v = element.get(qn(attr))
    return v if v is not None else default


def get_run_font_info(run):
    """提取 run 的字体信息 → dict"""
    rPr = run._element.find(qn('w:rPr'))
    info = {'cn': None, 'en': None, 'sz': None, 'bold': None, 'hint': None}
    if rPr is None:
        return info
    rf = rPr.find(qn('w:rFonts'))
    if rf is not None:
        info['cn'] = rf.get(qn('w:eastAsia'))
        info['en'] = rf.get(qn('w:ascii'))
        info['hint'] = rf.get(qn('w:hint'))
    sz_el = rPr.find(qn('w:sz'))
    if sz_el is not None:
        info['sz'] = int(sz_el.get(qn('w:val'))) // 2  # 半磅 → 磅
    b_el = rPr.find(qn('w:b'))
    info['bold'] = b_el is not None
    return info


def get_paragraph_format_info(p):
    """提取段落格式信息 → dict"""
    pPr = p._element.find(qn('w:pPr'))
    info = {
        'align': None, 'left_indent': None, 'right_indent': None,
        'first_line_indent': None, 'line_spacing': None, 'line_rule': None,
        'space_before': None, 'space_after': None, 'outline_lvl': None,
        'text': p.text,
    }
    if pPr is None:
        return info

    # 对齐
    jc = pPr.find(qn('w:jc'))
    if jc is not None:
        info['align'] = jc.get(qn('w:val'))

    # 缩进
    ind = pPr.find(qn('w:ind'))
    if ind is not None:
        for attr, key in [('w:left', 'left_indent'), ('w:right', 'right_indent'),
                          ('w:firstLine', 'first_line_indent')]:
            v = ind.get(qn(attr))
            if v is not None:
                info[key] = int(v) / 20  # twips → pt

    # 行距
    sp = pPr.find(qn('w:spacing'))
    if sp is not None:
        line = sp.get(qn('w:line'))
        rule = sp.get(qn('w:lineRule'))
        if line is not None:
            info['line_spacing'] = int(line)
        if rule is not None:
            info['line_rule'] = rule
        for attr, key in [('w:before', 'space_before'), ('w:after', 'space_after')]:
            v = sp.get(qn(attr))
            if v is not None:
                info[key] = int(v) / 20  # twips → pt (approx)

    # 大纲级别
    ol = pPr.find(qn('w:outlineLvl'))
    if ol is not None:
        info['outline_lvl'] = int(ol.get(qn('w:val')))

    return info


def get_paragraph_runs_info(p):
    """提取段落中所有 run 的字体信息（取第一个非空为准）"""
    runs_info = []
    for r in p.runs:
        info = get_run_font_info(r)
        info['text'] = r.text
        runs_info.append(info)
    return runs_info


def extract_page_setup(doc):
    """提取页面设置"""
    sec = doc.sections[0] if doc.sections else None
    if sec is None:
        return {}
    return {
        'page_width_cm': round(sec.page_width / 360000, 2),
        'page_height_cm': round(sec.page_height / 360000, 2),
        'margin_top_cm': round(sec.top_margin / 360000, 2),
        'margin_bottom_cm': round(sec.bottom_margin / 360000, 2),
        'margin_left_cm': round(sec.left_margin / 360000, 2),
        'margin_right_cm': round(sec.right_margin / 360000, 2),
    }


def _get_xml_style_info(doc, style_id):
    """从 styles.xml 直接读取样式精确值（绕过 python-docx 主题映射）"""
    try:
        sty = doc.styles[style_id]
    except KeyError:
        return None
    # 获取 styles.xml 根
    from docx.oxml.ns import qn as _qn
    styles_el = sty.element.getroottree().getroot()
    # 按 styleId 查找
    target = doc.styles[style_id].element

    info = {'name': style_id, 'font': {}, 'paragraph': {}}
    rPr = target.find(_qn('w:rPr'))
    pPr = target.find(_qn('w:pPr'))

    if rPr is not None:
        rf = rPr.find(_qn('w:rFonts'))
        if rf is not None:
            for attr, key in [('w:eastAsia', 'eastAsia'), ('w:ascii', 'ascii'),
                              ('w:hAnsi', 'hAnsi')]:
                v = rf.get(_qn(attr))
                if v:
                    info['font'][key] = v
        sz_el = rPr.find(_qn('w:sz'))
        if sz_el is not None:
            info['font']['size_half_pt'] = int(sz_el.get(_qn('w:val')))
            info['font']['size'] = info['font']['size_half_pt'] / 2
        info['font']['bold'] = rPr.find(_qn('w:b')) is not None

    if pPr is not None:
        jc = pPr.find(_qn('w:jc'))
        if jc is not None:
            info['paragraph']['alignment'] = jc.get(_qn('w:val'))
        sp = pPr.find(_qn('w:spacing'))
        if sp is not None:
            line = sp.get(_qn('w:line'))
            rule = sp.get(_qn('w:lineRule'))
            before = sp.get(_qn('w:before'))
            after = sp.get(_qn('w:after'))
            if line:
                info['paragraph']['line'] = int(line)
                info['paragraph']['line_rule'] = rule or 'auto'
                info['paragraph']['line_multiplier'] = round(int(line) / 240, 2)
            if before:
                info['paragraph']['space_before_pt'] = int(before) / 20
            if after:
                info['paragraph']['space_after_pt'] = int(after) / 20
        ol = pPr.find(_qn('w:outlineLvl'))
        if ol is not None:
            info['paragraph']['outline_lvl'] = int(ol.get(_qn('w:val')))

    # 补充：从 docDefaults 读取全局默认字体
    defaults = styles_el.find(_qn('w:docDefaults'))
    if defaults is not None:
        rPrDefault = defaults.find(_qn('w:rPrDefault'))
        if rPrDefault is not None:
            drPr = rPrDefault.find(_qn('w:rPr'))
            if drPr is not None:
                drf = drPr.find(_qn('w:rFonts'))
                if drf is not None:
                    info['doc_defaults'] = {}
                    for attr, key in [('w:eastAsia', 'eastAsia'), ('w:ascii', 'ascii'),
                                      ('w:hAnsi', 'hAnsi')]:
                        v = drf.get(_qn(attr))
                        if v:
                            info['doc_defaults'][key] = v
    return info


def extract_style_def(doc, style_name):
    """提取命名样式定义（先尝试 XML 精确读取，回退到 python-docx API）"""
    info = _get_xml_style_info(doc, style_name)
    if info:
        return info
    # 回退：python-docx
    try:
        sty = doc.styles[style_name]
    except KeyError:
        return None
    info = {'name': style_name, 'font': {}, 'paragraph': {}}
    f = sty.font
    info['font']['name'] = f.name
    info['font']['size'] = f.size.pt if f.size else None
    info['font']['bold'] = f.bold
    pf = sty.paragraph_format
    info['paragraph']['alignment'] = str(pf.alignment) if pf.alignment else None
    info['paragraph']['first_line_indent'] = pf.first_line_indent.pt if pf.first_line_indent else 0
    info['paragraph']['line_spacing'] = pf.line_spacing
    rPr = sty.element.rPr
    if rPr is not None:
        rf = rPr.find(qn('w:rFonts'))
        if rf is not None:
            info['font']['eastAsia'] = rf.get(qn('w:eastAsia'))
            info['font']['ascii'] = rf.get(qn('w:ascii'))
            info['font']['hAnsi'] = rf.get(qn('w:hAnsi'))
    return info


# ======================================================================
# 第三部分：辅助工具
# ======================================================================

def pt_to_zh(pt):
    """磅值 → 中文字号"""
    mapping = [
        (42, '初号'), (36, '小初'), (26, '一号'), (24, '小一'),
        (22, '二号'), (18, '小二'), (16, '三号'), (15, '小三'),
        (14, '四号'), (12, '小四'), (10.5, '五号'), (9, '小五'),
    ]
    for val, name in mapping:
        if abs(pt - val) < 0.5:
            return f'{name}（{pt}pt）'
    return f'{pt}pt'


def alignment_to_text(align_val):
    """对齐方式 → 中文描述"""
    m = {
        'left': '左对齐', 'right': '右对齐', 'center': '居中',
        'both': '两端对齐', 'justify': '两端对齐',
    }
    return m.get(str(align_val).lower(), str(align_val))


def spacing_to_text(rule, line_val):
    """行距值 → 中文描述"""
    if rule == 'auto':
        ratio = int(line_val) / 240
        return f'{ratio:.1f} 倍行距' if ratio != 1 else '单倍行距'
    else:
        return f'固定值 {int(line_val)/20:.0f} 磅'


def save_json(obj, path, indent=2):
    """保存 JSON 文件（自动处理中文引号等特殊字符）"""
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(obj, f, ensure_ascii=False, indent=indent)
    return path


def load_json(path):
    """加载 JSON 文件"""
    with open(path, 'r', encoding='utf-8') as f:
        return json.load(f)
    with open(path, 'r', encoding='utf-8') as f:
        return json.load(f)


def write_file(path, text):
    """写入文本文件"""
    with open(path, 'w', encoding='utf-8') as f:
        f.write(text)
    return path


def read_file(path):
    """读取文本文件"""
    with open(path, 'r', encoding='utf-8') as f:
        return f.read()


def find_vision_js():
    """查找 vision.js 脚本位置"""
    candidates = [
        os.path.expanduser('~/.claude/skills/vision-skill/vision.js'),
        os.path.expanduser('~/.claude/skills/vision-skill/vision.mjs'),
    ]
    for p in candidates:
        if os.path.exists(p):
            return p
    return None
