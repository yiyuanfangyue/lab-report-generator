# -*- coding: utf-8 -*-
"""
阶段 1：模板格式分析
====================
解析 .docx 实验报告模板，提取封面格式和正文样式。

交互流程：
  1. 提取页面设置、封面格式、正文样式
  2. 展示给用户确认，允许修改字号/字体/行距等
  3. 输出：
     格式规范.md       — 人类可读的 Markdown
     格式规范-LaTeX模板.tex — LaTeX 格式定义
     格式规范.json      — 机器可读格式数据（给阶段 3 用）

用法:
  python analyze_template.py <模板.docx> [输出前缀]
"""
import sys, os, re
from docx import Document
from docx.oxml.ns import qn
from lib.docx_utils import (
    extract_page_setup, extract_style_def,
    get_paragraph_format_info, get_paragraph_runs_info,
    save_json, write_file, pt_to_zh, alignment_to_text,
    spacing_to_text, get_run_font_info,
)


def extract_cover_paragraphs(doc):
    """提取封面段落列表（到第一个带有标题样式的段落为止）"""
    paragraphs = []
    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ''
        if 'Heading' in style_name or 'heading' in style_name:
            break
        fmt = get_paragraph_format_info(p)
        runs = get_paragraph_runs_info(p)
        drawing = p._element.findall('.//' + qn('w:drawing'))
        has_image = bool(drawing) or bool(p._element.findall('.//' + qn('w:object')))
        paragraphs.append({
            'text': p.text,
            'format': fmt,
            'runs': runs,
            'has_image': has_image,
            'style': style_name,
        })
    if len(paragraphs) == len(list(doc.paragraphs)):
        paragraphs = paragraphs[:15]
    return paragraphs


def _detect_cover(paragraphs):
    """
    通过启发式规则判断模板是否有封面页。
    封面页的典型特征：
    - 至少有一段居中（align=center）
    - 字号明显大于正文（>= 18pt）
    - 或者包含图片（校徽）
    - 或者文本含有典型封面关键词
    返回 True/False
    """
    if not paragraphs:
        return False

    # 封面特有的必填字段（这些基本只出现在封面）
    cover_only_fields = ['姓  名', '姓 名', '学  号', '学 号',
                         '班  级', '班 级', '学生姓名', '学生学号',
                         '学生班级', '专  业', '专 业', '指导教师',
                         '成  绩', '成 绩']

    # 封面常见关键词（可出现多次）
    cover_keywords = ['大学', '学院', '实验报告', '课程设计',
                      '实验名称', '实验课程']

    # 首段是否有图片
    first_has_image = paragraphs[0].get('has_image', False) if paragraphs else False

    # 统计特征
    all_text = ' '.join(p.get('text', '') for p in paragraphs)
    centered = sum(1 for p in paragraphs if p.get('format', {}).get('align') == 'center')
    has_image = any(p.get('has_image') for p in paragraphs)
    large_fonts = 0
    for p in paragraphs:
        for r in p.get('runs', []):
            sz = r.get('sz', 0) or 0
            if sz >= 26:  # >= 13pt
                large_fonts += 1
                break

    hits_cover_only = sum(1 for kw in cover_only_fields if kw in all_text)
    hits_keywords = sum(1 for kw in cover_keywords if kw in all_text)

    # 判定逻辑
    # 强信号：有封面专属字段（姓名/学号/班级等），几乎100%是封面
    if hits_cover_only >= 1:
        return True

    # 中强信号：首段有图片 + 有封面关键词（如"大学""实验报告"）
    if first_has_image and hits_keywords >= 1:
        return True

    # 弱信号组合：以下满足任意2条
    score = sum([
        centered >= 3,
        has_image and hits_keywords >= 1,
        large_fonts >= 2,
    ])
    return score >= 2


def extract_body_format(doc):
    """提取正文格式定义（从 XML 精确读取）"""
    return {
        'normal': extract_style_def(doc, 'Normal'),
        'heading1': extract_style_def(doc, 'Heading 1'),
        'heading2': extract_style_def(doc, 'Heading 2'),
    }


def print_format_summary(data):
    """打印格式摘要供用户确认"""
    ps = data['page_setup']
    print(f'\n  页面设置:')
    print(f'    纸张: A4 ({ps["page_width_cm"]}x{ps["page_height_cm"]}cm)')
    print(f'    页边距: 上{ps["margin_top_cm"]}cm 下{ps["margin_bottom_cm"]}cm 左{ps["margin_left_cm"]}cm 右{ps["margin_right_cm"]}cm')

    print(f'\n  封面: {"有（自动检测到）" if data.get("has_cover") else "无（自动判定为无封面）"}')
    if data.get('has_cover'):
        print(f'  封面段落: {len(data["cover_paragraphs"])} 段')
        for i, cp in enumerate(data['cover_paragraphs']):
            t = cp['text'][:40] or '(空/图)'
            print(f'    P{i}: {t}')

    bf = data['body_format']
    for key, label in [('heading1', '一级标题'), ('heading2', '二级标题'), ('normal', '正文')]:
        s = bf.get(key)
        if not s:
            continue
        f = s.get('font', {})
        cn = f.get('eastAsia') or (s.get('doc_defaults') or {}).get('eastAsia', '宋体')
        en = f.get('ascii') or (s.get('doc_defaults') or {}).get('ascii', 'Times New Roman')
        sz = f.get('size') or (f.get('size_half_pt', 24) / 2) or 12
        pf = s.get('paragraph', {})
        line = pf.get('line_multiplier', '')
        before = pf.get('space_before_pt', '')
        after = pf.get('space_after_pt', '')
        print(f'  {label}: {cn}/{en} {sz:.0f}pt 加粗={f.get("bold","")} '
              f'行距={line} 段前={before}pt 段后={after}pt')


def interactive_adjust(data):
    """交互式调整：让用户修改格式值"""
    print('\n' + '=' * 60)
    print('  请确认以下格式信息，可直接修改字号（其他项如需修改可手动编辑配置文件）')
    print('=' * 60)
    print_format_summary(data)

    # 让用户确认/修正封面判定
    detected = data.get('has_cover', False)
    label = '有' if detected else '无'
    ans = input(f'  封面判定 ({label})，是否正确？(y=正确 n=翻转 回车不变): ').strip().lower()
    if ans == 'n':
        data['has_cover'] = not detected
        if data['has_cover']:
            # 重新提取封面段落（从 doc 已经存了，需恢复到 data 中）
            print('  [注意] 请在后续手动编辑格式规范.json 补充封面段落内容')
        else:
            data['cover_paragraphs'] = []
        print(f'  已翻转为: {"有封面" if data["has_cover"] else "无封面"}')

    bf = data['body_format']
    for key, label in [('heading1', '一级标题'), ('heading2', '二级标题'), ('normal', '正文')]:
        s = bf.get(key)
        if not s:
            continue
        f = s['font']
        cur = f.get('size') or (f.get('size_half_pt', 24) / 2) or 12
        ans = input(f'  {label} 字号 (当前 {cur:.0f}pt, 回车不变): ').strip()
        if ans:
            try:
                new_sz = int(ans)
                f['size'] = new_sz
                f['size_half_pt'] = new_sz * 2
            except ValueError:
                pass

    ans = input('\n  格式确认无误？(y=继续, n=放弃): ').strip().lower()
    if ans == 'n':
        print('  已放弃')
        return False
    return True


def generate_format_md(page_setup, cover_paragraphs, body_format):
    """生成 Markdown 格式规范"""
    lines = []
    lines.append('# 实验报告格式规范（自动提取）')
    lines.append('')
    lines.append('## 页面设置')
    lines.append('')
    lines.append('| 项目 | 值 |')
    lines.append('|------|-----|')
    lines.append(f'| 纸张 | A4（{page_setup["page_width_cm"]} cm x {page_setup["page_height_cm"]} cm） |')
    lines.append(f'| 上边距 | {page_setup["margin_top_cm"]} cm |')
    lines.append(f'| 下边距 | {page_setup["margin_bottom_cm"]} cm |')
    lines.append(f'| 左边距 | {page_setup["margin_left_cm"]} cm |')
    lines.append(f'| 右边距 | {page_setup["margin_right_cm"]} cm |')
    lines.append('')
    lines.append('---')
    lines.append('')

    if cover_paragraphs:
        lines.append('## 封面格式')
        lines.append('')
        for i, cp in enumerate(cover_paragraphs):
            lines.append(f'### P{i}: {cp["text"][:50] or "(空行/图片)"}')
            lines.append('')
            if cp['has_image']:
                lines.append('- 类型：图片（嵌入）')
            fmt = cp['format']
            if fmt['align']:
                lines.append(f'- 对齐：{alignment_to_text(fmt["align"])}')
            if fmt['first_line_indent']:
                lines.append(f'- 首行缩进：{fmt["first_line_indent"]:.1f} pt')
            if fmt['left_indent']:
                lines.append(f'- 左缩进：{fmt["left_indent"]:.1f} pt')
            if fmt['right_indent']:
                lines.append(f'- 右缩进：{fmt["right_indent"]:.1f} pt')
            if fmt['line_spacing']:
                lines.append(f'- 行距：{spacing_to_text(fmt["line_rule"], fmt["line_spacing"])}')
            for ri, r in enumerate(cp['runs']):
                if not r.get('text', '').strip() and cp.get('has_image'):
                    continue
                if not r.get('text', '').strip() and r.get('cn') is None:
                    continue
                parts = []
                if r.get('cn'):
                    parts.append(f'中文字体：{r["cn"]}')
                if r.get('en'):
                    parts.append(f'英文字体：{r["en"]}')
                if r.get('sz'):
                    parts.append(f'字号：{r["sz"]}pt（{pt_to_zh(r["sz"])}）')
                if r.get('bold'):
                    parts.append('加粗')
                if parts:
                    lines.append(f'- 字体：{"，".join(parts)}')
            lines.append('')

    lines.append('---')
    lines.append('')
    lines.append('## 正文样式')
    lines.append('')

    for style_name in ['heading1', 'heading2', 'normal']:
        style = body_format.get(style_name)
        if not style:
            continue
        label = {'heading1': '一级标题', 'heading2': '二级标题', 'normal': '正文（Normal）'}[style_name]
        lines.append(f'### {label}')
        lines.append('')
        lines.append('| 属性 | 值 |')
        lines.append('|------|-----|')
        f = style.get('font', {})
        cn = f.get('eastAsia') or (style.get('doc_defaults') or {}).get('eastAsia', '宋体')
        en = f.get('ascii') or (style.get('doc_defaults') or {}).get('ascii', 'Times New Roman')
        sz = f.get('size') or (f.get('size_half_pt', 24) / 2) or 12
        lines.append(f'| 中文字体 | {cn} |')
        lines.append(f'| 英文字体 | {en} |')
        lines.append(f'| 字号 | {sz:.0f} pt（{pt_to_zh(sz)}） |')
        lines.append(f'| 加粗 | {"是" if f.get("bold") else "否"} |')
        pf = style.get('paragraph', {})
        if pf.get('alignment'):
            lines.append(f'| 对齐 | {alignment_to_text(pf["alignment"])} |')
        if pf.get('space_before_pt'):
            lines.append(f'| 段前 | {pf["space_before_pt"]:.0f} pt |')
        if pf.get('space_after_pt'):
            lines.append(f'| 段后 | {pf["space_after_pt"]:.0f} pt |')
        if pf.get('line_multiplier'):
            lines.append(f'| 行距 | {pf["line_multiplier"]} 倍行距 |')
        lines.append('')

    return '\n'.join(lines)


def generate_latex_template(page_setup, cover_paragraphs, body_format):
    """生成 LaTeX 格式模板"""
    lines = []
    mt = page_setup['margin_top_cm']
    mb = page_setup['margin_bottom_cm']
    ml = page_setup['margin_left_cm']
    mr = page_setup['margin_right_cm']

    lines.append('% !TEX program = xelatex')
    lines.append('% 格式模板 — 由 lab-report-generator 自动生成')
    lines.append(r'\documentclass[12pt,a4paper,fontset=fandol]{ctexart}')
    lines.append('')
    lines.append(r'\usepackage{geometry}')
    lines.append(f'\\geometry{{top={mt}cm,bottom={mb}cm,left={ml}cm,right={mr}cm,headheight=1.5cm,footskip=1.7cm}}')
    lines.append('')
    lines.append(r'\usepackage{setspace,graphicx,caption,fancyhdr,ulem,enumitem,float,booktabs,array}')
    lines.append('')

    normal = body_format.get('normal', {})
    nf = normal.get('font', {})
    n_size = nf.get('size') or nf.get('size_half_pt', 24) / 2 or 12
    lines.append(f'% 正文：{n_size:.0f}pt，首行缩进2字符')
    lines.append(r'\setlength{\parindent}{' + f'{int(n_size*2)}pt' + '}')
    lines.append(r'\setlength{\parskip}{0pt}')
    lines.append('')

    h1 = body_format.get('heading1', {})
    h2 = body_format.get('heading2', {})
    h1f = h1.get('font', {}) if h1 else {}
    h2f = h2.get('font', {}) if h2 else {}
    h1p = h1.get('paragraph', {}) if h1 else {}
    h2p = h2.get('paragraph', {}) if h2 else {}

    h1_size = int(h1f.get('size') or h1f.get('size_half_pt', 36) / 2 or 18)
    h2_size = int(h2f.get('size') or h2f.get('size_half_pt', 32) / 2 or 16)
    h1_before = int(h1p.get('space_before_pt') or h1p.get('space_before', 17))
    h1_after = int(h1p.get('space_after_pt') or h1p.get('space_after', 16.5))
    h2_before = int(h2p.get('space_before_pt') or h2p.get('space_before', 13))
    h2_after = int(h2p.get('space_after_pt') or h2p.get('space_after', 13))

    zh_sizes = {42: '0', 36: '-0', 26: '1', 24: '-1', 22: '2', 18: '-2', 16: '3', 15: '-3', 14: '4', 12: '-4'}
    h1_zh = zh_sizes.get(h1_size, '-2')
    h2_zh = zh_sizes.get(h2_size, '3')

    lines.append(r'\ctexset{')
    lines.append('  section = {')
    lines.append('    format     = \\zihao{' + str(h1_zh) + r'}\bfseries\setstretch{1.5},')
    lines.append('    name       = {},')
    lines.append('    number     = \\chinese{section},')
    lines.append('    aftername  = {、},')
    lines.append(f'    beforeskip = {h1_before}pt,')
    lines.append(f'    afterskip  = {h1_after}pt,')
    lines.append('    indent     = 0pt,')
    lines.append('  },')
    lines.append('  subsection = {')
    lines.append('    format     = \\zihao{' + str(h2_zh) + r'}\bfseries\setstretch{1.5},')
    lines.append('    name       = {},')
    lines.append(f'    number     = \\arabic{{section}}.\\arabic{{subsection}},')
    lines.append('    aftername  = \\hspace{1em},')
    lines.append(f'    beforeskip = {h2_before}pt,')
    lines.append(f'    afterskip  = {h2_after}pt,')
    lines.append('    indent     = 0pt,')
    lines.append('  },')
    lines.append('}')
    lines.append('')

    lines.append(r'\pagestyle{fancy}')
    lines.append(r'\fancyhf{}')
    lines.append(r'\fancyfoot[C]{\thepage}')
    lines.append(r'\renewcommand{\headrulewidth}{0pt}')
    lines.append('')
    lines.append(r'\graphicspath{{./}}')
    lines.append(r'\DeclareGraphicsExtensions{.jpg,.png,.pdf}')
    lines.append('')

    if cover_paragraphs:
        lines.append(r'\begin{document}')
        lines.append(r'\begin{titlepage}')
        lines.append(r'  \setlength{\parindent}{0pt}')
        lines.append(r'  \setstretch{1.5}')
        lines.append('')
        for i, cp in enumerate(cover_paragraphs):
            text = cp['text']
            fmt = cp['format']
            runs = cp['runs']
            if cp['has_image']:
                lines.append(f'  % P{i}: [图片]')
                lines.append(r'  \begin{center}')
                lines.append(r'    \includegraphics[width=5cm]{logo.jpeg}% 替换为实际图片')
                lines.append(r'  \end{center}')
                lines.append('')
                continue
            if not text.strip():
                lines.append(f'  % P{i}: 空行')
                lines.append(r'  \vspace{1cm}')
                lines.append('')
                continue
            font_info = runs[0] if runs else {}
            sz = font_info.get('sz', 12)
            bold = font_info.get('bold', False)
            zh_sz = zh_sizes.get(sz, '-4')
            bf_cmd = r'\bfseries' if bold else ''
            lines.append(f'  % P{i}: {text[:50]}')
            li = fmt.get('left_indent', 0)
            ri = fmt.get('right_indent', 0)
            fi = fmt.get('first_line_indent', 0)
            if li or ri:
                if li:
                    lines.append(r'  \leftskip=' + f'{li:.1f}pt')
                if ri:
                    lines.append(r'  \rightskip=' + f'{ri:.1f}pt')
                lines.append(f'  {{\\zihao{{{zh_sz}}}{bf_cmd} {text}\\par}}')
                if li:
                    lines.append(r'  \leftskip=0pt')
                if ri:
                    lines.append(r'  \rightskip=0pt')
            elif fi:
                lines.append(r'  \setlength{\parindent}{' + f'{fi:.1f}pt' + '}')
                lines.append(f'  {{\\zihao{{{zh_sz}}}{bf_cmd} {text}\\par}}')
                lines.append(r'  \setlength{\parindent}{0pt}')
            else:
                lines.append(r'  \begin{center}')
                lines.append(f'    {{\\zihao{{{zh_sz}}}{bf_cmd} {text}\\par}}')
                lines.append(r'  \end{center}')
            lines.append('')
        lines.append(r'\end{titlepage}')
        lines.append('')
    else:
        lines.append(r'\begin{document}')
        lines.append('')

    lines.append(r'% ===== 正文开始 =====')
    lines.append(r'\setlength{\parindent}{' + f'{int(n_size*2)}pt' + '}')
    lines.append(r'\setstretch{1.667}')
    lines.append(r'\setcounter{page}{1}')
    lines.append('')
    lines.append(r'% 在此处编写正文内容...')
    lines.append('')
    lines.append(r'\section{实验目的}')
    lines.append(r'\section{实验原理}')
    lines.append(r'\section{实验步骤}')
    lines.append(r'\section{实验作业}')
    lines.append(r'\subsection{作业题1}')
    lines.append(r'\subsection{作业题2}')
    lines.append('')
    lines.append(r'\end{document}')

    return '\n'.join(lines)


def extract_template_format(template_path):
    """主函数：分析模板"""
    doc = Document(template_path)
    cover_paragraphs = extract_cover_paragraphs(doc)
    has_cover = _detect_cover(cover_paragraphs)
    return {
        'has_cover': has_cover,
        'page_setup': extract_page_setup(doc),
        'cover_paragraphs': cover_paragraphs if has_cover else [],
        'body_format': extract_body_format(doc),
    }


# ======================================================================
if __name__ == '__main__':
    # 检查 --yes 标志
    auto_confirm = '--yes' in sys.argv
    args = [a for a in sys.argv[1:] if a != '--yes']

    if len(args) < 1:
        print('用法: python analyze_template.py [--yes] <模板.docx> [输出前缀]')
        print('说明: 分析 .docx 模板格式 -> 格式规范.md + LaTeX模板.tex + 格式规范.json')
        print('  --yes  自动模式，跳过交互式字号调整，直接输出')
        sys.exit(1)

    template_path = args[0]
    if not os.path.exists(template_path):
        print(f'错误: 文件不存在 — {template_path}')
        sys.exit(1)

    prefix = args[1] if len(args) > 1 else '格式规范'
    output_md = f'{prefix}.md'
    output_tex = f'{prefix}-LaTeX模板.tex'
    output_json = f'{prefix}.json'

    data = extract_template_format(template_path)

    # 交互确认
    ok = True
    if not auto_confirm:
        ok = interactive_adjust(data)
    if not ok:
        sys.exit(1)

    # 生成输出文件
    md_content = generate_format_md(
        data['page_setup'], data['cover_paragraphs'], data['body_format'])
    tex_content = generate_latex_template(
        data['page_setup'], data['cover_paragraphs'], data['body_format'])

    write_file(output_md, md_content)
    write_file(output_tex, tex_content)
    save_json(data, output_json)

    print(f'\n[OK] 模板格式分析完成')
    print(f'  封面段落: {len(data["cover_paragraphs"])}')
    print(f'  输出:')
    print(f'    {output_md}')
    print(f'    {output_tex}')
    print(f'    {output_json}')
