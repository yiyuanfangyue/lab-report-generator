# -*- coding: utf-8 -*-
"""
阶段 3：报告生成
================
将前面各阶段的产出合并，生成完整的实验报告。

步骤：
  3a. 分析模板 .docx 的内容结构（作为内容骨架参考）
  3b. 合并 → 完整 LaTeX 文档
  3c. 用户选择 PDF / DOCX，编译输出

用法:
  python generate_report.py --content 实验内容结构.json --latex-format LaTeX格式模板.tex
                           --images image_map.json [--template 模板.docx]
                           [--output 输出前缀] [--format pdf|docx]
"""
import sys, os, json, subprocess, shutil, re, argparse
from lib.docx_utils import (
    save_json, load_json, write_file, read_file,
    set_xml_font, set_xml_indent, set_xml_spacing,
    load_img_bytes, cv_line, cv_empty,
)
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn


# ╔══════════════════════════════════════════════════════════════════════╗
# ║ 步骤 3a: 模板内容结构分析                                           ║
# ╚══════════════════════════════════════════════════════════════════════╝

def analyze_template_content(template_path):
    """
    读取模板 .docx 的正文内容结构。
    返回段落列表，每段含文本和样式名。
    封面段落自动跳过（复用阶段1的分析结果）。
    """
    if not template_path or not os.path.exists(template_path):
        return None

    from docx import Document
    from lib.docx_utils import get_paragraph_format_info

    doc = Document(template_path)
    sections = []
    in_body = False

    for p in doc.paragraphs:
        text = p.text.strip()
        style = p.style.name if p.style else ''
        fmt = get_paragraph_format_info(p)

        # 检测正文入口：有 Heading 样式或 outlineLvl
        pPr = p._element.find(qn('w:pPr'))
        if pPr is not None:
            ol = pPr.find(qn('w:outlineLvl'))
            sty_el = pPr.find(qn('w:pStyle'))
            is_heading = False
            if ol is not None:
                is_heading = True
            if sty_el is not None:
                sval = sty_el.get(qn('w:val'))
                if sval and 'Heading' in sval:
                    is_heading = True
            if is_heading:
                in_body = True

        if not in_body and not text:
            continue
        if not in_body:
            continue

        # 检测页面分隔符
        pb = p._element.findall(qn('w:lastRenderedPageBreak'))
        if pb:
            sections.append({'type': 'page_break'})

        entry = {'type': 'body', 'text': text, 'style': style}
        if 'Heading 1' in style:
            entry['type'] = 'heading1'
        elif 'Heading 2' in style:
            entry['type'] = 'heading2'
        elif 'Heading' in style:
            entry['type'] = 'heading'
        sections.append(entry)

    return sections if sections else None


# ╔══════════════════════════════════════════════════════════════════════╗
# ║ 步骤 3b: LaTeX 文档生成                                            ║
# ╚══════════════════════════════════════════════════════════════════════╝

def _escape_latex(text):
    """转义 LaTeX 特殊字符"""
    replacements = [
        ('\\', '\\textbackslash{}'),
        ('{', '\\{'), ('}', '\\}'),
        ('$', '\\$'), ('&', '\\&'), ('#', '\\#'),
        ('^', '\\^{}'), ('_', '\\_'), ('%', '\\%'),
        ('~', '\\textasciitilde{}'),
    ]
    for old, new in replacements:
        text = text.replace(old, new)
    return text


def _zh_section_number(num):
    """阿拉伯数字 → 中文数字（用于 \chinese{section}）"""
    # LaTeX 的 \chinese 命令会自动处理，这里仅用于展示/调试
    return str(num)


def build_latex_content(section, image_map, level=1):
    """生成单个章节的 LaTeX 代码"""
    t = section['type']
    lines = []

    if t == 'heading1':
        title = section.get('title', '')
        # 移除已有的中文编号前缀，让 LaTeX 自动编号
        clean = re.sub(r'^[一二三四五六七八九十]+[、.．]\s*', '', title)
        lines.append(f'\\section{{{_escape_latex(clean)}}}')
        lines.append('')

    elif t == 'heading2':
        title = section.get('title', '')
        clean = re.sub(r'^[\d]+[、.．]\s*', '', title)
        lines.append(f'\\subsection{{{_escape_latex(clean)}}}')
        lines.append('')

    elif t == 'heading3':
        title = section.get('title', '')
        lines.append(f'\\subsubsection*{{{_escape_latex(title)}}}')
        lines.append('')

    elif t == 'body':
        text = section.get('text', '')
        if text:
            lines.append(_escape_latex(text) + '\\par')
        else:
            lines.append('\\vspace{1em}')

    elif t == 'list_item':
        text = section.get('text', '')
        lines.append(_escape_latex(text) + '\\par')

    elif t == 'image_placeholder':
        img_id = section.get('id', '')
        caption = section.get('caption', '')
        instruction = section.get('text', '')
        raw_paths = image_map.get(img_id, '')
        # 兼容单路径（字符串）和多路径（列表）
        if isinstance(raw_paths, str):
            paths = [raw_paths] if raw_paths else []
        else:
            paths = raw_paths or []

        # 输出指令说明文字
        if instruction and instruction.strip():
            lines.append(_escape_latex(instruction.strip()) + '\\par')
            lines.append('')

        if paths:
            # 生成有意义的图注
            short_caps = _get_short_captions(img_id, paths, instruction.strip(), caption)
            for idx, p in enumerate(paths):
                if not p or not os.path.exists(p):
                    continue
                rel = os.path.basename(p)
                rel_safe = rel.replace(' ', '_').replace('（', '(').replace('）', ')')
                rel_safe = os.path.splitext(rel_safe)[0] + '.jpg'
                cap_text = short_caps[idx] if idx < len(short_caps) else (caption or instruction.strip()[:60])
                lines.append(r'\begin{figure}[H]')
                lines.append(r'  \centering')
                lines.append(f'  \\includegraphics[width=0.85\\textwidth]{{{rel_safe}}}')
                if cap_text:
                    lines.append(f'  \\caption{{{_escape_latex(cap_text)}}}')
                lines.append(r'\end{figure}')
                lines.append('')
        else:
            lines.append(f'[{img_id}] {_escape_latex(instruction)}\\par')
            lines.append('')

    # 处理子章节
    children = section.get('children', [])
    for child in children:
        child_lines = build_latex_content(child, image_map, level + 1)
        lines.extend(child_lines)

    return lines


def generate_latex(latex_format_path, content_json_path, image_map_path, output_path=None, template_path=None):
    """生成完整 LaTeX 文档"""
    # -- 读取输入 --
    content = load_json(content_json_path)
    image_map_data = load_json(image_map_path)

    # 提取 image_map（兼容两种格式）
    if 'image_map' in image_map_data:
        image_map = image_map_data['image_map']
    else:
        image_map = image_map_data

    # -- 读取 LaTeX 格式模板开头 --
    if latex_format_path and os.path.exists(latex_format_path):
        latex_preamble = read_file(latex_format_path)
    else:
        latex_preamble = _default_latex_preamble()

    # -- 找到正文开始标记 --
    doc_begin = latex_preamble.find(r'\begin{document}')
    if doc_begin >= 0:
        preamble = latex_preamble[:doc_begin]
        after_begin = latex_preamble[doc_begin:]
    else:
        preamble = latex_preamble
        after_begin = r'\begin{document}'

    # -- 从 after_begin 中去除模板自带的 \end{document} 和占位内容 --
    # 只保留到 \end{titlepage} 为止（封面部分），正文内容由下面代码生成
    end_doc = after_begin.find(r'\end{document}')
    if end_doc >= 0:
        after_begin = after_begin[:end_doc]
    # 去掉模板中占位的章节示例（"在此处编写正文内容"之后的内容）
    placeholder_marker = after_begin.find(r'在此处编写正文内容')
    if placeholder_marker >= 0:
        after_begin = after_begin[:placeholder_marker]

    # 如果模板无封面，去掉 \begin{titlepage}...\end{titlepage} 块
    # 读取格式数据判断 has_cover
    has_cover_latex = True
    fmt_path = content_json_path.replace('实验内容结构.json', '格式规范.json')
    if os.path.exists(fmt_path):
        try:
            fmt_data = load_json(fmt_path)
            has_cover_latex = fmt_data.get('has_cover', True)
        except Exception:
            pass
    if not has_cover_latex:
        # 去掉 titlepage 环境
        tp_begin = after_begin.find(r'\begin{titlepage}')
        tp_end = after_begin.find(r'\end{titlepage}')
        if tp_begin >= 0 and tp_end > tp_begin:
            # 保留 \begin{document} 和 titlepage 之后的内容
            before_tp = after_begin[:tp_begin]
            after_tp = after_begin[tp_end + len(r'\end{titlepage}'):]
            after_begin = before_tp + after_tp
            print('  模板无封面，已移除 LaTeX titlepage 环境')

    # -- 构建正文 --
    body_lines = []
    body_lines.append(r'\setlength{\parindent}{21pt}')
    body_lines.append(r'\setstretch{1.667}')
    body_lines.append(r'\setcounter{page}{1}')
    body_lines.append('')

    sections = content.get('sections', [])
    for sec in sections:
        sec_lines = build_latex_content(sec, image_map)
        body_lines.extend(sec_lines)

    # -- 组装 --
    full = []
    full.append(preamble)
    full.append(after_begin)
    full.extend(body_lines)
    full.append(r'\end{document}')

    tex_content = '\n'.join(full)

    # -- 输出 --
    if not output_path:
        output_path = '完整实验报告.tex'
    write_file(output_path, tex_content)

    # -- 复制图片到 LaTeX 目录（重命名：空格 → 下划线） --
    tex_dir = os.path.dirname(os.path.abspath(output_path))
    copied = 0

    # 从模板 .docx 提取校徽
    if template_path and os.path.exists(template_path):
        try:
            import zipfile
            z = zipfile.ZipFile(template_path)
            for name in z.namelist():
                if 'media' in name and name.lower().endswith(('.jpeg', '.jpg')):
                    badge_data = z.read(name)
                    badge_dest = os.path.join(tex_dir, 'logo.jpeg')
                    with open(badge_dest, 'wb') as f:
                        f.write(badge_data)
                    copied += 1
                    break
        except Exception:
            pass
    for img_id, raw_paths in image_map.items():
        if isinstance(raw_paths, str):
            paths = [raw_paths] if raw_paths else []
        else:
            paths = raw_paths or []
        for p in paths:
            if p and os.path.exists(p):
                fname = os.path.basename(p)
                safe_name = fname.replace(' ', '_').replace('（', '(').replace('）', ')')
                dest = os.path.join(tex_dir, safe_name)
                if os.path.abspath(p) != os.path.abspath(dest):
                    # 复制时检查 RGBA → RGB，并另存为 JPEG（xelatex 对 JPEG 兼容性更好）
                    from PIL import Image as _PilImage
                    _img = _PilImage.open(p)
                    if _img.mode == 'RGBA':
                        _bg = _PilImage.new('RGB', _img.size, (255, 255, 255))
                        _bg.paste(_img, mask=_img.split()[3])
                        _img = _bg
                    # 确保最大宽度不超过 2000px
                    if _img.size[0] > 2000:
                        _ratio = 2000 / _img.size[0]
                        _img = _img.resize((2000, int(_img.size[1] * _ratio)), _PilImage.LANCZOS)
                    dest_jpeg = os.path.splitext(dest)[0] + '.jpg'
                    _img.save(dest_jpeg, 'JPEG', quality=85, optimize=True)
                    copied += 1

    print(f'[OK] 完整 LaTeX 文档已生成: {output_path}')
    print(f'  已复制 {copied} 张图片到 LaTeX 目录')
    return output_path


def _default_latex_preamble():
    """默认 LaTeX 格式模板（当无模板分析结果时使用）"""
    return r"""% !TEX program = xelatex
% 默认 LaTeX 报告格式模板
\documentclass[12pt,a4paper,fontset=fandol]{ctexart}

\usepackage{geometry}
\geometry{top=2.54cm,bottom=2.54cm,left=3.18cm,right=3.18cm,headheight=1.5cm,footskip=1.7cm}

\usepackage{setspace,graphicx,caption,fancyhdr,ulem,enumitem,float,booktabs,array}
\usepackage[bf,labelsep=space]{caption}

\setlength{\parindent}{24pt}
\setlength{\parskip}{0pt}

\ctexset{
  section = {
    format     = \zihao{-2}\bfseries\setstretch{1.5},
    name       = {},
    number     = \chinese{section},
    aftername  = {、},
    beforeskip = 17pt,
    afterskip  = 16.5pt,
    indent     = 0pt,
  },
  subsection = {
    format     = \zihao{3}\bfseries\setstretch{1.5},
    name       = {},
    number     = \arabic{section}.\arabic{subsection},
    aftername  = \hspace{1em},
    beforeskip = 13pt,
    afterskip  = 13pt,
    indent     = 0pt,
  },
}

\pagestyle{fancy}
\fancyhf{}
\fancyfoot[C]{\thepage}
\renewcommand{\headrulewidth}{0pt}

\graphicspath{{./}}
\DeclareGraphicsExtensions{.jpg,.png,.pdf}

\begin{document}
"""


# ╔══════════════════════════════════════════════════════════════════════╗
# ║ 步骤 3c-1: 编译 PDF（xelatex）                                     ║
# ╚══════════════════════════════════════════════════════════════════════╝

def compile_pdf(tex_path):
    """编译 PDF：优先 lualatex（兼容性好），回退到 xelatex"""
    tex_dir = os.path.dirname(os.path.abspath(tex_path))
    tex_name = os.path.basename(tex_path)

    compiler = None
    for c in ['lualatex', 'xelatex']:
        if shutil.which(c):
            compiler = c
            break

    if not compiler:
        print()
        print('=' * 60)
        print('  [!!] 未检测到 LaTeX 编译器')
        print('=' * 60)
        print('  生成 PDF 需要安装 MiKTeX 或 TeX Live。')
        print()
        print('  安装方法：')
        print('    Windows: https://miktex.org/download  (下载安装包安装)')
        print('    其他:    sudo apt install texlive-xetex 或 brew install texlive')
        print()
        print('  安装后需确保 lualatex 或 xelatex 命令可在终端中运行。')
        print()
        print('  如果不方便安装，请使用 DOCX 格式代替 PDF。')
        print('=' * 60)
        print()
        return None

    print(f'  使用 {compiler} 编译...')
    for i in range(2):
        print(f'  第{i+1}次...')
        subprocess.run(
            [compiler, '-interaction=nonstopmode', tex_name],
            cwd=tex_dir, capture_output=True, timeout=180,
        )

    pdf_path = os.path.splitext(tex_path)[0] + '.pdf'
    if os.path.exists(pdf_path):
        size_kb = os.path.getsize(pdf_path) // 1024
        print(f'[OK] PDF 编译成功: {pdf_path} ({size_kb} KB)')
        return pdf_path
    else:
        print(f'[!] PDF 未生成，请检查 LaTeX 错误')
        log_path = os.path.splitext(tex_path)[0] + '.log'
        if os.path.exists(log_path):
            with open(log_path, 'r', encoding='utf-8', errors='ignore') as f:
                for l in f.readlines()[-30:]:
                    print(f'  {l.rstrip()[:120]}')
        return None


# ╔══════════════════════════════════════════════════════════════════════╗
# ║ 步骤 3c-2: 生成 DOCX（python-docx）                                ║
# ╚══════════════════════════════════════════════════════════════════════╝

def _get_short_captions(img_id, paths, instruction, caption=''):
    """为图片生成简短有意义的图注

    自动从 content JSON 的 image_placeholder 字段读取图注文字：
    - 单图（1 张）：直接用 caption 或 instruction 文字
    - 多图（N 张）：用 caption / instruction 为基础，自动追加编号 (1/N) (2/N) ...
    - 不再使用硬编码映射，不同实验的截图由用户在 content JSON 中自定义 text/caption
    """
    if not paths:
        return []

    # 选择最佳基础文字：caption > instruction > fallback
    base = (caption or instruction or f'图 {img_id}').strip()
    # 限制长度
    if len(base) > 60:
        base = base[:57] + '...'

    n = len(paths)
    if n == 1:
        return [base]
    else:
        return [f'{base}（{i+1}/{n}）' for i in range(n)]


def build_docx_cover(doc, cover_data, template_path=None):
    """根据封面格式数据重建封面"""
    if not cover_data:
        return False

    # 从模板 .docx 提取校徽图片到临时文件
    badge_file = None
    if template_path and os.path.exists(template_path):
        try:
            import zipfile, io, tempfile
            z = zipfile.ZipFile(template_path)
            for name in z.namelist():
                if 'media' in name and name.lower().endswith(('.jpeg', '.jpg')):
                    badge_data = z.read(name)
                    temp = tempfile.NamedTemporaryFile(delete=False, suffix='.jpeg')
                    temp.write(badge_data)
                    temp.close()
                    badge_file = temp.name
                    break
        except Exception:
            pass

    for cp in cover_data:
        text = cp.get('text', '')
        fmt = cp.get('format', {})
        runs = cp.get('runs', [])
        has_image = cp.get('has_image', False)

        if has_image:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            set_xml_spacing(p, 360, False)
            if badge_file and os.path.exists(badge_file):
                p.add_run().add_picture(badge_file, width=Cm(5.19))
            p.add_run('')
            continue

        fi = fmt.get('first_line_indent', 0)
        li = fmt.get('left_indent', 0)
        ri = fmt.get('right_indent', 0)

        if not text.strip():
            cv_empty(doc)
            continue

        # 提取第一个 run 的字体
        cn = '宋体'
        sz = 12
        bold = False
        if runs:
            r0 = runs[0]
            cn = r0.get('cn') or '宋体'
            sz = r0.get('sz') or 12
            bold = r0.get('bold', False)

        # 对齐：有左/右缩进时左对齐，有首行缩进时默认左对齐，否则居中
        align = WD_ALIGN_PARAGRAPH.CENTER
        if li or ri or fi:
            align = WD_ALIGN_PARAGRAPH.LEFT
        if fmt.get('align') == 'left':
            align = WD_ALIGN_PARAGRAPH.LEFT
        elif fmt.get('align') == 'right':
            align = WD_ALIGN_PARAGRAPH.RIGHT

        # P7: 实验名称 — 用分段 run 实现下划线效果（值部分加下划线）
        if sz == 18 and '实验名称' in text and li is None and fi:
            # 拆分成 "实验名称：" 和值 两部分
            p = doc.add_paragraph()
            p.alignment = align
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            set_xml_spacing(p, 360, False)
            if fi:
                set_xml_indent(p, first=fi)
            label_text = '实验名称：'
            value_text = text.replace('实验名称：', '').strip()
            r1 = p.add_run(label_text)
            set_xml_font(r1, '楷体_GB2312', sz=sz, bold=True)
            r2 = p.add_run(value_text)
            set_xml_font(r2, '黑体', sz=sz, bold=False, underline='single')
        else:
            cv_line(doc, text, sz, bold=bold, align=align,
                    li=li if li else None, ri=ri if ri else None,
                    fi=fi if fi else None, cn=cn)

    # 清理临时文件
    if badge_file and os.path.exists(badge_file):
        try: os.unlink(badge_file)
        except: pass

    return True


def build_docx_body(doc, sections, image_map, format_config=None):
    """根据内容结构和格式配置生成正文"""
    # 默认格式
    h1_font = {'cn': '宋体', 'en': None, 'sz': 18, 'bold': True}
    h2_font = {'cn': '宋体', 'en': None, 'sz': 16, 'bold': True}
    body_font = {'cn': '宋体', 'en': 'Times New Roman', 'sz': 12, 'bold': False}

    # 正文段落格式（Normal）
    body_first_indent = 21  # pt
    body_line_rule = 'exact'  # 固定值
    body_line_val = 400  # 20pt = 400 half-pts

    # 从格式配置中读取（若有）
    if format_config:
        bf = format_config.get('body_format', {})
        h1s = bf.get('heading1', {})
        h2s = bf.get('heading2', {})
        ns = bf.get('normal', {})
        def _get_sz(d, default):
            f = d.get('font', {})
            return int(f.get('size') or f.get('size_half_pt', default * 2) / 2 or default)

        h1_font['sz'] = _get_sz(h1s, 18)
        h2_font['sz'] = _get_sz(h2s, 16)
        body_font['sz'] = _get_sz(ns, 12)

        # 读取正文段落格式（行距、缩进）
        npf = ns.get('paragraph', {})
        if npf.get('line') and npf.get('line_rule'):
            body_line_val = int(npf['line'])
            body_line_rule = npf['line_rule']
        # 首行缩进 = 字号 * 2
        body_first_indent = int(body_font['sz'] * 2)

        # 字体继承链
        def _get_cn(d):
            f = d.get('font', {})
            cn = f.get('eastAsia')
            if cn: return cn
            dd = d.get('doc_defaults') or {}
            return dd.get('eastAsia', '宋体')
        def _get_en(d):
            f = d.get('font', {})
            en = f.get('ascii')
            if en: return en
            dd = d.get('doc_defaults') or {}
            return dd.get('ascii', 'Times New Roman')

        h1_font['cn'] = _get_cn(h1s)
        h2_font['cn'] = _get_cn(h2s)
        body_font['cn'] = _get_cn(ns)
        h1_font['en'] = _get_en(h1s)
        h2_font['en'] = _get_en(h2s)
        body_font['en'] = _get_en(ns)

    def _write_heading1(title):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(17)
        p.paragraph_format.space_after = Pt(16.5)
        set_xml_spacing(p, 360, False)
        r = p.add_run(title)
        set_xml_font(r, h1_font['cn'], h1_font.get('en'), h1_font['sz'], h1_font['bold'])
        return p

    def _write_heading2(title):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(13)
        p.paragraph_format.space_after = Pt(13)
        set_xml_spacing(p, 360, False)
        r = p.add_run(title)
        set_xml_font(r, h2_font['cn'], h2_font.get('en'), h2_font['sz'], h2_font['bold'])
        return p

    # 图片计数器（自动标号 图1、图2...）
    fig_counter = [0]  # 用列表实现闭包可写

    def _write_body(text):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(body_first_indent)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        set_xml_spacing(p, body_line_val, str(body_line_rule).startswith('exact'))
        r = p.add_run(text)
        set_xml_font(r, body_font['cn'], body_font.get('en'), body_font['sz'], body_font['bold'])
        return p

    def _write_fig(img_path, caption, w_cm=14.5):
        nonlocal fig_counter
        # 读取图片格式配置，没有则用默认值
        img_fmt = {}
        cap_fmt = {}
        if format_config:
            img_fmt = format_config.get('image_format', {}).get('paragraph', {})
            cap_fmt = format_config.get('image_format', {}).get('caption', {})

        buf = load_img_bytes(img_path, 1200)
        if buf:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(img_fmt.get('space_before_pt', 6))
            p.paragraph_format.space_after = Pt(img_fmt.get('space_after_pt', 2))
            # 图片段落必须用自动行距
            set_xml_spacing(p, 360, False)
            p.add_run().add_picture(buf, width=Cm(w_cm))
        if caption:
            fig_counter[0] += 1
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.first_line_indent = Pt(0)
            p2.paragraph_format.space_before = Pt(2)
            p2.paragraph_format.space_after = Pt(6)
            # 自动标号 图N:
            label = f'图{fig_counter[0]} {caption}'
            r = p2.add_run(label)
            cn = cap_fmt.get('cn') or body_font['cn']
            en = cap_fmt.get('en') or body_font.get('en')
            sz = cap_fmt.get('sz') or body_font['sz']
            bold = cap_fmt.get('bold', True)
            set_xml_font(r, cn, en, sz, bold)

    def _process_section(sec):
        t = sec.get('type', '')
        children = sec.get('children', [])

        if t == 'heading1':
            title = sec.get('title', '')
            _write_heading1(title)
        elif t == 'heading2':
            title = sec.get('title', '')
            _write_heading2(title)
        elif t == 'heading3':
            title = sec.get('title', '')
            # heading3 用二级标题格式（同 heading2）
            _write_heading2(title)
        elif t == 'body':
            text = sec.get('text', '')
            if text:
                _write_body(text)
        elif t == 'list_item':
            text = sec.get('text', '')
            if text:
                _write_body(text)  # 列表项也按正文格式
        elif t == 'image_placeholder':
            img_id = sec.get('id', '')
            caption = sec.get('caption', '')
            instruction = sec.get('text', '')
            raw_paths = image_map.get(img_id, '')
            if isinstance(raw_paths, str):
                paths = [raw_paths] if raw_paths else []
            else:
                paths = raw_paths or []

            # 先输出指令说明文字
            if instruction and instruction.strip():
                _write_body(instruction.strip())

            # 生成有意义的图注：优先用 caption 字段，其次 instruction，最后 fallback
            base_ins = (caption or instruction.strip() or f'图{img_id}').strip()[:80]
            # 对多图场景生成简短图注
            short_caps = _get_short_captions(img_id, paths, base_ins, caption)
            for idx, p in enumerate(paths):
                if not p or not os.path.exists(p):
                    continue
                cap = short_caps[idx] if idx < len(short_caps) else base_ins[:60]
                _write_fig(p, cap)

        for child in children:
            _process_section(child)

    for sec in sections:
        _process_section(sec)


def build_docx(content_json_path, format_data_path, image_map_path, template_path,
               output_path=None):
    """生成完整 DOCX 文档"""
    # -- 读取输入 --
    content = load_json(content_json_path)
    image_map_data = load_json(image_map_path)

    if 'image_map' in image_map_data:
        image_map = image_map_data['image_map']
    else:
        image_map = image_map_data

    # 读取封面格式数据
    cover_data = None
    body_format = None
    page_setup = None
    has_cover = True  # 默认有封面（向后兼容）
    if format_data_path and os.path.exists(format_data_path):
        fmt = load_json(format_data_path)
        cover_data = fmt.get('cover_paragraphs')
        body_format = fmt.get('body_format')
        page_setup = fmt.get('page_setup')
        has_cover = fmt.get('has_cover', True)

    # 或者从模板直接读取
    if not cover_data and template_path and os.path.exists(template_path):
        try:
            from analyze_template import extract_template_format
            data = extract_template_format(template_path)
            cover_data = data['cover_paragraphs']
            body_format = data['body_format']
            page_setup = data['page_setup']
            has_cover = data.get('has_cover', True)
        except Exception as e:
            print(f'  [!] 读取模板格式失败: {e}')

    # -- 创建文档 --
    doc = Document()

    # 页面设置
    if page_setup:
        sec = doc.sections[0]
        try:
            sec.page_width = Cm(page_setup.get('page_width_cm', 21.0))
            sec.page_height = Cm(page_setup.get('page_height_cm', 29.7))
            sec.top_margin = Cm(page_setup.get('margin_top_cm', 2.54))
            sec.bottom_margin = Cm(page_setup.get('margin_bottom_cm', 2.54))
            sec.left_margin = Cm(page_setup.get('margin_left_cm', 3.18))
            sec.right_margin = Cm(page_setup.get('margin_right_cm', 3.18))
        except Exception:
            pass

    # -- Normal 样式 --
    sty = doc.styles['Normal']
    sty.font.name = 'Times New Roman'
    sty.font.size = Pt(12)
    sty.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # -- 封面（如果有） --
    if has_cover and cover_data:
        build_docx_cover(doc, cover_data, template_path)
        doc.add_page_break()
    elif has_cover:
        print('  [注意] 模板判定为有封面但未提取到封面数据')
    else:
        print('  模板无封面，跳过封面生成')

    # -- 正文 --
    format_config = {'body_format': body_format} if body_format else None
    sections = content.get('sections', [])
    build_docx_body(doc, sections, image_map, format_config)

    # -- 保存 --
    if not output_path:
        output_path = '完整实验报告.docx'

    doc.save(output_path)
    size_kb = os.path.getsize(output_path) // 1024
    print(f'[OK] DOCX 已生成: {output_path} ({size_kb} KB)')
    return output_path


# ╔══════════════════════════════════════════════════════════════════════╗
# ║ MD 预览生成                                                         ║
# ╚══════════════════════════════════════════════════════════════════════╝

def generate_md_preview(content_json_path, image_map_path, output_path='预览_实验报告.md'):
    """生成 Markdown 预览文档（含结构、图片占位）"""
    content = load_json(content_json_path)
    image_data = load_json(image_map_path)
    image_map = image_data.get('image_map', image_data)

    lines = []
    lines.append('# 实验报告预览')
    lines.append('')
    lines.append('> 请确认以下内容结构是否正确。图片以 `[图片占位]` 标记。')
    lines.append('')

    # 封面信息
    cover = content.get('cover', {})
    cover_kids = cover.get('children', [])
    if cover_kids:
        lines.append('---')
        lines.append('## 封面')
        lines.append('')
        for c in cover_kids:
            t = c.get('text', '')
            if c.get('type') == 'image':
                lines.append('`[校徽图片]`')
                lines.append('')
            elif t.strip():
                lines.append(t)
                lines.append('')
        lines.append('---')
        lines.append('')

    # 正文
    for sec in content.get('sections', []):
        _md_render_node(sec, image_map, lines, 0)

    write_file(output_path, '\n'.join(lines))
    print(f'[OK] MD 预览已生成: {output_path}')
    return output_path


def _md_render_node(node, image_map, lines, depth):
    """递归渲染节点为 MD"""
    prefix = '  ' * depth
    t = node.get('type', '')
    title = node.get('title', node.get('text', ''))
    children = node.get('children', [])

    if t == 'heading1':
        lines.append(f'{prefix}## {title}')
        lines.append('')
    elif t == 'heading2':
        lines.append(f'{prefix}### {title}')
        lines.append('')
    elif t == 'heading3':
        lines.append(f'{prefix}#### {title}')
        lines.append('')
    elif t == 'body':
        if title.strip():
            lines.append(f'{prefix}{title}')
            lines.append('')
    elif t == 'image_placeholder':
        img_id = node.get('id', '')
        ins_text = node.get('text', '').strip()
        if ins_text:
            lines.append(f'{prefix}{ins_text}')
            lines.append('')
        raw = image_map.get(img_id, '')
        paths = raw if isinstance(raw, list) else ([raw] if raw else [])
        count = len([p for p in paths if p and os.path.exists(p)])
        if count > 0:
            lines.append(f'{prefix}`[图片: {img_id} — {count}张截图]`')
        else:
            lines.append(f'{prefix}`[图片: {img_id} — 待插入{node.get("expected_count",1)}张截图]`')
        lines.append('')

    for child in children:
        _md_render_node(child, image_map, lines, depth + 1)

def main():
    parser = argparse.ArgumentParser(description='阶段3: 实验报告生成器')
    parser.add_argument('--content', required=True, help='实验内容结构.json')
    parser.add_argument('--latex-format', default='格式规范-LaTeX模板.tex',
                        help='LaTeX格式模板.tex')
    parser.add_argument('--images', required=True, help='image_map.json')
    parser.add_argument('--template', help='原始模板.docx（可选，用于内容结构参考）')
    parser.add_argument('--format-data', help='格式数据.json（来自阶段1的机器可读数据）')
    parser.add_argument('--output', default='完整实验报告', help='输出文件前缀')
    parser.add_argument('--format', choices=['pdf', 'docx', 'both', 'md', 'all'], default=None,
                        help='输出格式（默认先问用户）')
    args = parser.parse_args()

    # -- 检查文件 --
    for f in [args.content, args.images]:
        if not os.path.exists(f):
            print(f'错误: 文件不存在 — {f}')
            sys.exit(1)

    # -- 校验 image_map 中的图片路径 --
    try:
        img_data = load_json(args.images)
        img_map = img_data.get('image_map', img_data)
        missing = []
        empty = []
        for img_id, raw_paths in img_map.items():
            if isinstance(raw_paths, str):
                if raw_paths:
                    paths = [raw_paths]
                else:
                    paths = []
                    empty.append(img_id)
            elif isinstance(raw_paths, list):
                if raw_paths:
                    paths = raw_paths
                else:
                    paths = []
                    empty.append(img_id)
            else:
                paths = raw_paths or []
            for p in paths:
                if not p:
                    empty.append(img_id)
                elif not os.path.exists(p):
                    missing.append((img_id, p))
        if empty:
            print(f'[!] 警告: 以下插入点尚未映射图片（路径为空）: {", ".join(empty)}')
        if missing:
            print(f'[!] 警告: 以下图片文件不存在:')
            for iid, p in missing:
                print(f'      [{iid}] {p}')
            print('  请先修正 image_map.json 中的路径')
    except Exception as e:
        print(f'[!] 无法读取图片映射文件: {e}')

    latex_format = args.latex_format
    if not os.path.exists(latex_format):
        print(f'[!] 未找到 LaTeX 格式模板 {latex_format}，将使用默认格式')
        latex_format = None

    # -- 步骤 3a: 分析模板内容结构（日志输出，暂不深度使用） --
    if args.template and os.path.exists(args.template):
        print('步骤 3a: 分析模板内容结构...')
        template_content = analyze_template_content(args.template)
        if template_content:
            print(f'  模板正文段落数: {len(template_content)}')
        else:
            print('  模板内容分析未返回数据')
    else:
        print('步骤 3a: 未提供模板，跳过内容结构分析')

    # -- 步骤 3b: 生成 LaTeX --
    print('\n步骤 3b: 生成完整 LaTeX 文档...')
    tex_path = generate_latex(latex_format, args.content, args.images,
                              f'{args.output}.tex', template_path=args.template)

    # -- 步骤 3c: 选择输出格式 --
    output_format = args.format
    if not output_format:
        print('\n步骤 3c: 选择输出格式')
        print('  [1] PDF（通过 xelatex 编译）')
        print('  [2] DOCX（通过 python-docx 生成）')
        print('  [3] 两者都生成')
        print('  [4] 先生成 MD 预览给我确认')
        choice = input('  请选择 (1/2/3/4): ').strip()
        output_format = {'1': 'pdf', '2': 'docx', '3': 'both', '4': 'md'}.get(choice, 'pdf')

    results = []

    # 先生成 MD 预览（如果选 md 或 all）
    if output_format in ('md', 'all'):
        print('\n--- 生成 MD 预览 ---')
        md_path = generate_md_preview(args.content, args.images, f'{args.output}.md')
        results.append(md_path)
        print('\n请查看并确认 MD 预览内容，如需修改请告知。确认后重新运行指定最终格式。')
        if output_format == 'md':
            return

    if output_format in ('pdf', 'both', 'all'):
        print('\n--- 生成 PDF ---')
        pdf_path = compile_pdf(tex_path)
        if pdf_path:
            results.append(pdf_path)

    if output_format in ('docx', 'both', 'all'):
        print('\n--- 生成 DOCX ---')
        docx_path = build_docx(args.content, args.format_data, args.images,
                               args.template, f'{args.output}.docx')
        results.append(docx_path)

    print(f'\n{"="*50}')
    if results:
        print('[OK] 报告生成完成!')
        for r in results:
            print(f'  {r}')
    else:
        print('[!] 报告生成失败')
        sys.exit(1)


if __name__ == '__main__':
    main()
